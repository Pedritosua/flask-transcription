import openai
from flask import Flask, request, render_template, send_file, redirect, url_for
import os
from pydub import AudioSegment
import speech_recognition as sr
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
TRANSCRIPTS_FOLDER = 'transcripts'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TRANSCRIPTS_FOLDER, exist_ok=True)

# Configura tu clave API directamente en el código
openai.api_key = "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"  # Reemplaza esto con tu clave real


def transcribir_audio_en_segmentos(filepath, segment_duration=30):
    recognizer = sr.Recognizer()
    audio = AudioSegment.from_file(filepath)
    duration = len(audio) / 1000  # Duración en segundos
    transcripcion_completa = ""

    for inicio in range(0, int(duration), segment_duration):
        fin = min(inicio + segment_duration, duration)
        segmento = audio[inicio * 1000:fin * 1000]
        segmento_wav_path = "temp_segment.wav"
        segmento.export(segmento_wav_path, format="wav")

        try:
            with sr.AudioFile(segmento_wav_path) as source:
                audio_data = recognizer.record(source)
                texto = recognizer.recognize_google(audio_data, language="es-ES")
                transcripcion_completa += texto + " "
        except sr.UnknownValueError:
            transcripcion_completa += "[No se pudo entender el audio] "
        except sr.RequestError as e:
            return f"Error al conectar con el servicio de reconocimiento de voz; {e}"
        finally:
            if os.path.exists(segmento_wav_path):
                os.remove(segmento_wav_path)

    return transcripcion_completa


def generar_acta_con_ia(transcripcion):
    # Define el contexto del asistente para generar el acta
    messages = [
        {"role": "system", "content": "Eres un asistente experto en generar actas formales para reuniones."},
        {"role": "user", "content": f"""
        Transforma la siguiente transcripción en un acta formal con el formato típico que incluye:
        - Encabezado con fecha y lugar de la reunión.
        - Lista de asistentes.
        - Temas tratados.
        - Decisiones tomadas y tareas asignadas.

        Transcripción:
        {transcripcion}
        """}
    ]

    try:
        # Llamada al modelo de ChatCompletion
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",  # Usa gpt-3.5-turbo si no tienes acceso a gpt-4
            messages=messages,
            temperature=0.7,
            max_tokens=1000
        )
        # Retorna el contenido generado por la IA
        return response['choices'][0]['message']['content'].strip()
    except openai.error.AuthenticationError as e:
        return f"Error de autenticación: {str(e)}. Verifica tu clave API."
    except openai.error.OpenAIError as e:
        return f"Error al conectar con OpenAI: {str(e)}."


@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return "No se seleccionó un archivo"
        file = request.files['file']
        if file.filename == '':
            return "Archivo vacío"
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)

            # Convertir el archivo a .wav
            wav_path = os.path.join(UPLOAD_FOLDER, 'temp_audio.wav')
            audio = AudioSegment.from_file(filepath)
            audio.export(wav_path, format="wav")

            # Transcribir el audio en segmentos
            transcripcion = transcribir_audio_en_segmentos(wav_path)

            if "Error" in transcripcion:
                return transcripcion

            # Guardar la transcripción en un archivo de Word
            word_path = os.path.join(TRANSCRIPTS_FOLDER, 'transcripcion.docx')
            document = Document()
            document.add_paragraph(transcripcion)
            document.save(word_path)

            # Eliminar archivos temporales
            os.remove(filepath)
            os.remove(wav_path)

            return redirect(url_for('transcription_done'))

    return render_template('upload.html')


@app.route('/done')
def transcription_done():
    return """
    <h1>Hecho, ya puedes descargar tu archivo</h1>
    <a href="/download/transcripcion.docx" style="margin-right: 20px;">Descargar Transcripción</a>
    <button onclick="window.location.href='/generate-acta'" style="padding: 10px; background-color: #007BFF; color: white; border: none; border-radius: 5px; cursor: pointer; margin-right: 20px;">
        Generar Acta
    </button>
    <button onclick="window.location.href='/'" style="padding: 10px; background-color: #28A745; color: white; border: none; border-radius: 5px; cursor: pointer;">
        Cargar un nuevo audio
    </button>
    """


@app.route('/generate-acta', methods=['GET'])
def generate_acta():
    # Lee el archivo de transcripción generado
    filepath = os.path.join(TRANSCRIPTS_FOLDER, 'transcripcion.docx')
    document = Document(filepath)
    transcripcion = "\n".join([p.text for p in document.paragraphs])

    # Genera el acta usando la IA
    acta = generar_acta_con_ia(transcripcion)

    # Guarda el acta en un archivo Word
    acta_path = os.path.join(TRANSCRIPTS_FOLDER, 'acta.docx')
    doc = Document()
    doc.add_paragraph(acta)
    doc.save(acta_path)

    # Muestra la opción de descargar el acta
    return """
    <h1>Acta generada correctamente</h1>
    <a href="/download/acta.docx" style="margin-right: 20px;">Descargar Acta</a>
    <button onclick="window.location.href='/'" style="padding: 10px; background-color: #007BFF; color: white; border: none; border-radius: 5px; cursor: pointer;">
        Cargar un nuevo audio
    </button>
    """


@app.route('/download/<filename>')
def download_file(filename):
    path = os.path.join(TRANSCRIPTS_FOLDER, filename)
    return send_file(path, as_attachment=True, download_name=filename)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)
